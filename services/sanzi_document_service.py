# Source Generated with Decompyle++
# File: sanzi_document_service.pyc (Python 3.11)

from __future__ import annotations
import csv
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import date
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from services.output_trace_service import add_trace_to_report
from services.sanzi_rules import expected_filename, land_actuality_label, land_status_label, proof_accessory_type, proof_material_name, requirements_for_status, use_status_label
_INVALID_FILENAME = re.compile('[<>:"/\\\\|?*\\x00-\\x1f]')
TEMPLATE_FILENAME = '三资情况证明模板.docx'
FLY_LAND_TEMPLATE_FILENAME = '外部（内部）飞地证明模板.docx'
VILLAGE_COMMITTEE_TEMPLATE_FILENAME = '争议待确权地证明模板.docx'
BUNDLED_TEMPLATE_DIRNAME = '三资使用状态模板'
SUPPORTED_PLACEHOLDERS = ('{{行政村}}', '{{图斑编号}}', '{{乡镇}}', '{{面积（亩）}}', '{{图上面积（亩）}}', '{{使用状态}}', '{{地类现状}}', '{{使用人}}', '{{身份证号}}', '{{备注}}', '{{所属组}}', '{{发包人}}', '{{承包人}}', '{{企业名称}}', '{{工作进度}}', '{{行政区编码}}', '{{当前日期}}')

def _safe_text(value = None, default = None):
    if value in (None, '', [], { }):
        return default
    text = str(value).strip()
    return default if text.lower() in frozenset({'nan', 'none', 'null', 'undefined'}) else text


def _first(record = None, *keys, default = None):
    for source in (record, record.get('detail') if isinstance(record, dict) else None, record.get('summary') if isinstance(record, dict) else None):
        if not isinstance(source, dict):
            continue
        for key in keys:
            value = _safe_text(source.get(key))
            if value:
                return value
    return default


def _number(record = None, *keys, default = None):
    value = _first(record, *keys)
    if value == '':
        return default
    try:
        return int(float(value))
    except Exception:
        return default


def normalize_land_record(record = None):
    landcode = _first(record, 'landcode', 'landCode')
    use_status = _number(record, 'usestatus', 'useStatus', default = -1)
    actuality = _number(record, 'landactuality', 'landActuality', default = -1)
    land_status = _number(record, 'landstatus', 'landStatus', default = -1)
    area = _first(record, 'land_area', 'area', 'mapArea')
    try:
        area = f'{float(area):.4f}'.rstrip('0').rstrip('.') if area else ''
    except Exception:
        pass

    employer = _first(record, 'employer', 'issuer', 'contractor', 'contractorName', 'userperson', 'useperson')
    person_name = _first(record, 'name', 'individualName', 'legalPersonName',
                         'username', 'userName', 'usagePerson', 'usePerson',
                         'personName', 'contractorName')
    company_name = _first(record, 'companyname', 'companyName')
    usage_person = person_name or company_name or employer
    return {
        'landcode': landcode,
        'town': _first(record, 'distictn_1', 'districtn_1', 'town', 'townname',
                       'townName', 'township', 'townshipName', 'streetName'),
        'village': _first(record, 'districtname', 'districtName', 'distinctName',
                          'village', 'villageName', 'domaindistrictname',
                          'domainDistrictName', 'village_name', 'orgName'),
        'districtcode': _first(record, 'districtcode', 'districtCode', 'distinctCode'),
        'landname': _first(record, 'landname', 'landName', 'plotName', default=''),
        'group': _first(record, 'ownershipgrup', 'ownershipGroup', 'group'),
        'area': area,
        'use_status': use_status,
        'use_status_label': use_status_label(use_status),
        'land_actuality': actuality,
        'land_actuality_label': land_actuality_label(actuality),
        'land_status': land_status,
        'land_status_label': land_status_label(land_status),
        'employer': employer,
        'person_name': person_name,
        'usage_person': usage_person,
        'idnumber': _first(record, 'idnumber', 'idNumber', 'identityNumber',
                           'identityCard', 'idcard', 'idCard', 'cardNo'),
        'companyname': company_name,
        'remark': _first(record, 'remark', 'remarks', 'memo'),
    }


def _set_run_font(run = None, size = None, bold = None, name = '宋体'):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold


def _resource_path(filename = None):
    candidates = []
    if getattr(sys, 'frozen', False):
        base = Path(getattr(sys, '_MEIPASS', Path(sys.executable).parent))
        candidates.extend([
            base / 'resources' / filename,
            Path(sys.executable).parent / 'resources' / filename])
    root = Path(__file__).resolve().parent.parent
    candidates.append(root / 'resources' / filename)
    for p in candidates:
        if p.exists() and p.is_file():
            return p
    return None


def _build_template(output_path = None, kind = None):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(16)
    titles = {
        'generic': '情  况  证  明',
        'fly': '飞 地 证 明',
        'village': '村 委 会 证 明' }
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(22)
    _set_run_font(title.add_run(titles.get(kind, titles['generic'])), 22, True, '方正小标宋简体')
    if kind == 'fly':
        body = '兹证明，{{乡镇}}{{行政村}}农村集体“三资”监管大数据平台图斑编号为{{图斑编号}}，面积为{{面积（亩）}}亩，使用状态为{{使用状态}}，地类现状为{{地类现状}}。经核实，该地块属于本村飞地，现使用人为{{使用人}}，身份证号为{{身份证号}}。备注：{{备注}}。'
    elif kind == 'village':
        body = '兹证明，{{乡镇}}{{行政村}}农村集体“三资”监管大数据平台图斑编号为{{图斑编号}}，面积为{{面积（亩）}}亩，使用状态为{{使用状态}}，地类现状为{{地类现状}}。该地块现使用人为{{使用人}}，身份证号为{{身份证号}}。经村委会调查核实，相关情况如下：{{备注}}。'
    else:
        body = '兹有{{乡镇}}{{行政村}}农村集体“三资”监管大数据平台图斑，图斑编号为{{图斑编号}}，面积为{{面积（亩）}}亩。经现场核实，该地块目前使用状态为{{使用状态}}，地类现状为{{地类现状}}，使用人为{{使用人}}，身份证号为{{身份证号}}。备注：{{备注}}。'
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(12)
    _set_run_font(p.add_run(body), 16)
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Pt(32)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _set_run_font(p2.add_run('以上情况属实，特此证明。'), 16)
    for _ in range(4):
        doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(sign.add_run('{{行政村}}股份经济合作社（盖章）'), 16)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(date_p.add_run('{{当前日期}}'), 16)
    output_path.parent.mkdir(parents = True, exist_ok = True)
    doc.save(output_path)
    return output_path


def create_default_template(output_path = None, *, overwrite, kind):
    output_path = Path(output_path)
    if output_path.exists() and not overwrite:
        return output_path
    output_path.parent.mkdir(parents = True, exist_ok = True)
    bundled = _resource_path(output_path.name)
    if bundled and bundled.resolve() != output_path.resolve():
        shutil.copy2(bundled, output_path)
        return output_path
    return _build_template(output_path, kind)


def _bundled_status_template_dir():
    candidates = []
    if getattr(sys, 'frozen', False):
        base = Path(getattr(sys, '_MEIPASS', Path(sys.executable).parent))
        candidates.extend([
            base / 'resources' / BUNDLED_TEMPLATE_DIRNAME,
            Path(sys.executable).parent / 'resources' / BUNDLED_TEMPLATE_DIRNAME,
            Path(sys.executable).parent / '_internal' / 'resources' / BUNDLED_TEMPLATE_DIRNAME])
    root = Path(__file__).resolve().parent.parent
    candidates.append(root / 'resources' / BUNDLED_TEMPLATE_DIRNAME)
    for folder in candidates:
        if folder.is_dir():
            return folder
    return None


def restore_default_templates(template_dir = None, *, overwrite):
    '''恢复12种内置模板，只覆盖同名默认模板，不删除用户自定义模板。'''
    folder = Path(template_dir)
    folder.mkdir(parents = True, exist_ok = True)
    restored = { }
    bundled_dir = _bundled_status_template_dir()
    if bundled_dir:
        for source in sorted(bundled_dir.glob('*.docx'), key = lambda x: x.name):
            target = folder / source.name
            if overwrite or not target.exists():
                shutil.copy2(source, target)
            restored[source.name] = target
    for filename, kind in (('三资情况证明模板.docx', 'generic'), ('村委会证明模板.docx', 'village')):
        target = folder / filename
        source = _resource_path(filename)
        if source and source.resolve() != target.resolve():
            if overwrite or not target.exists():
                shutil.copy2(source, target)
        elif overwrite or not target.exists():
            _build_template(target, kind)
        restored[filename] = target
    write_template_folder_guide(folder)
    return restored


def ensure_project_templates(project_dir = None):
    '''首次创建项目时补齐12种默认模板，已有同名文件不覆盖。'''
    template_dir = Path(project_dir) / '01_自定义模板'
    restored = restore_default_templates(template_dir, overwrite = False)
    copied = {k: v for k, v in restored.items()}
    copied['generic'] = template_dir / TEMPLATE_FILENAME
    return copied


def ensure_project_template(project_dir = None):
    return ensure_project_templates(project_dir)['generic']


def _normalize_template_name(value = None):
    return re.sub('[\\s_\\-—－（）()【】\\[\\]·.]+', '', str(value or '')).lower()


def list_template_files(template_dir = None):
    folder = Path(template_dir)
    if not folder.is_dir():
        return []
    return sorted([p for p in folder.glob('*.docx')], key = lambda x: x.name)


def _template_keywords_for_record(record = None):
    data = normalize_land_record(record)
    status = data['use_status']
    words = [data['use_status_label'], proof_material_name(status)]
    aliases = {
        -1: ('未填写', '默认', '通用'),
        0: ('对外发包', '发包'),
        2: ('征收征占', '征收', '征占'),
        3: ('闲置',),
        4: ('飞地', '飞地证明'),
        5: ('自留地', '菜地'),
        6: ('抵顶地',),
        7: ('经营性', '集体自用经营'),
        8: ('公共性', '集体自用公共'),
        9: ('争议待确权', '村委会证明'),
        10: ('田间硬化', '硬化路面'),
        11: ('村民自用',),
        12: ('无争议未确权',),
        13: ('延包后', '再分地'),
        14: ('已确权', '确权'),
    }
    words.extend(aliases.get(status, ()))
    return [w for w in words if w]


def choose_template_for_record(template_dir = None, record = None):
    '''按文件名自动选择模板。

    优先级：完整使用状态名称 > 状态关键词/材料名称 > 默认/通用模板 > 文件夹首个模板。
    用户只需把不同地块种类的 docx 放进同一个模板文件夹，并让文件名包含状态名称。
    '''
    templates = list_template_files(template_dir)
    if not templates:
        raise FileNotFoundError(f'模板文件夹中没有 Word 模板：{template_dir}')
    data = normalize_land_record(record)
    exact = _normalize_template_name(data['use_status_label'])
    keywords = _template_keywords_for_record(record)
    best = None
    for path in templates:
        name = _normalize_template_name(path.stem)
        score = 0
        if exact and exact != '未填写' and exact in name:
            score = 120
        else:
            for word in keywords:
                if word and word in name:
                    score = max(score, 80 + min(30, len(word)))
        if any(kw in name for kw in ('默认', '通用', '三资情况证明', '情况证明模板')):
            score = max(score, 20)
        candidate = (score, path.name.lower(), path)
        if best is None or candidate[0] > best[0] or (candidate[0] == best[0] and candidate[1] < best[1]):
            best = candidate
    assert best is not None
    return best[2]


def write_template_folder_guide(template_dir = None):
    folder = Path(template_dir)
    folder.mkdir(parents = True, exist_ok = True)
    readme = folder / '模板说明.txt'
    if readme.exists():
        return
    lines = [
        '===== 证明模板文件夹说明 =====',
        '',
        '本文件夹存放 Word (.docx) 格式的证明模板。',
        '支持的占位符（用双花括号包裹）：',
    ]
    for ph in SUPPORTED_PLACEHOLDERS:
        lines.append(f'  {ph}')
    lines.extend([
        '',
        '模板选择规则：',
        '  系统会根据图斑的"使用状态"自动选择文件名中包含对应状态名称的模板。',
        '  如果找不到匹配的模板，将使用文件名中包含"默认"或"通用"的模板。',
        '',
        '文件名示例：',
        '  对外发包模板.docx',
        '  飞地证明模板.docx',
        '  默认通用模板.docx',
    ])
    readme.write_text('\n'.join(lines), encoding = 'utf-8')


def _iter_paragraphs_in_container(container):
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_container(cell)


def _replace_placeholder_in_paragraph(paragraph = None, placeholder = None, value = None):
    full = ''.join(run.text for run in paragraph.runs)
    if placeholder not in full:
        return
    new_text = full.replace(placeholder, value)
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ''


def replace_placeholders(doc = None, replacements = None):
    for p in _iter_paragraphs_in_container(doc):
        for placeholder, value in replacements.items():
            _replace_placeholder_in_paragraph(p, placeholder, value)


def template_values(record = None):
    data = normalize_land_record(record)
    return {
        '{{行政村}}': data.get('village', ''),
        '{{图斑编号}}': data.get('landcode', ''),
        '{{乡镇}}': data.get('town', ''),
        '{{面积（亩）}}': data.get('area', ''),
        '{{图上面积（亩）}}': data.get('area', ''),
        '{{使用状态}}': data.get('use_status_label', ''),
        '{{地类现状}}': data.get('land_actuality_label', ''),
        '{{使用人}}': data.get('usage_person', ''),
        '{{身份证号}}': data.get('idnumber', ''),
        '{{备注}}': data.get('remark', ''),
        '{{所属组}}': data.get('group', ''),
        '{{发包人}}': data.get('employer', ''),
        '{{承包人}}': data.get('person_name', ''),
        '{{企业名称}}': data.get('companyname', ''),
        '{{工作进度}}': data.get('remark', ''),
        '{{行政区编码}}': data.get('districtcode', ''),
        '{{当前日期}}': date.today().strftime('%Y年%m月%d日'),
    }


def template_text(template_path = None):
    doc = Document(template_path)
    parts = []
    for p in _iter_paragraphs_in_container(doc):
        parts.append(p.text)
    return '\n'.join(parts)


def template_missing_placeholders(template_path = None):
    text = template_text(template_path)
    return [ph for ph in SUPPORTED_PLACEHOLDERS if ph not in text]


def fill_template(record = None, template_path = None, output_path = None):
    doc = Document(template_path)
    replacements = template_values(record)
    replace_placeholders(doc, replacements)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents = True, exist_ok = True)
    doc.save(output_path)
    return output_path


def _safe_folder_name(text = None):
    return _INVALID_FILENAME.sub('_', str(text or '')).strip()


def land_material_dir(project_dir = None, record = None):
    data = normalize_land_record(record)
    landcode = data.get('landcode') or '未知编号'
    village = data.get('village') or '未知村'
    folder = Path(project_dir) / '02_证明材料' / _safe_folder_name(village) / _safe_folder_name(landcode)
    folder.mkdir(parents = True, exist_ok = True)
    return folder


def _write_land_readme(record = None, output_dir = None):
    data = normalize_land_record(record)
    readme = Path(output_dir) / '图斑信息.txt'
    lines = [
        f'图斑编号：{data.get("landcode", "")}',
        f'行政村：{data.get("village", "")}',
        f'乡镇：{data.get("town", "")}',
        f'面积：{data.get("area", "")} 亩',
        f'使用状态：{data.get("use_status_label", "")}',
        f'地类现状：{data.get("land_actuality_label", "")}',
        f'使用人：{data.get("usage_person", "")}',
    ]
    readme.write_text('\n'.join(lines), encoding = 'utf-8')


def _write_manifest(project_dir = None, rows = None):
    manifest_dir = Path(project_dir) / '00_汇总'
    manifest_dir.mkdir(parents = True, exist_ok = True)
    csv_path = manifest_dir / '图斑清单.csv'
    with open(csv_path, 'w', newline = '', encoding = 'utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames = rows[0].keys() if rows else [])
        writer.writeheader()
        writer.writerows(rows)
    json_path = manifest_dir / '图斑清单.json'
    json_path.write_text(json.dumps(rows, ensure_ascii = False, indent = 2), encoding = 'utf-8')
    return csv_path, json_path


def _template_for_status(generic_template = None, use_status = None):
    return generic_template


def generate_project_documents(records = None, selected_landcodes = None, *,
                               overwrite = False, template_path = None,
                               template_dir = None, progress = None):
    records = list(records) if records else []
    rows = []
    results = []
    for record in records:
        data = normalize_land_record(record)
        landcode = data.get('landcode', '')
        if selected_landcodes and landcode not in selected_landcodes:
            continue
        tpl_dir = Path(template_dir) if template_dir else None
        if tpl_dir and tpl_dir.is_dir():
            tpl = choose_template_for_record(tpl_dir, record)
        elif template_path:
            tpl = Path(template_path)
        else:
            tpl = None
        project_dir = Path(template_dir).parent if template_dir else Path('.')
        out_dir = land_material_dir(project_dir, record)
        filename = expected_filename(data)
        docx_path = out_dir / filename
        if tpl and (overwrite or not docx_path.exists()):
            try:
                fill_template(record, tpl, docx_path)
            except Exception:
                pass
        _write_land_readme(record, out_dir)
        rows.append(data)
        results.append(docx_path)
        if progress:
            try:
                progress(landcode)
            except Exception:
                pass
    if rows:
        _write_manifest(Path(template_dir).parent if template_dir else Path('.'), rows)
    return results


def _find_office_converter():
    for name in ('soffice', 'libreoffice'):
        path = shutil.which(name)
        if path:
            return path
    return None


def _docx_to_pdf_word(docx_path = None, pdf_path = None):
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat = 17)
        doc.Close()
        word.Quit()
        return True, ''
    except Exception as e:
        return False, str(e)


def _docx_to_pdf_office(docx_path = None, pdf_path = None):
    converter = _find_office_converter()
    if not converter:
        return False, '未找到 LibreOffice / soffice'
    try:
        result = subprocess.run(
            [converter, '--headless', '--convert-to', 'pdf',
             '--outdir', str(pdf_path.parent), str(docx_path)],
            capture_output = True, text = True, timeout = 120)
        if result.returncode == 0:
            generated = pdf_path.parent / (docx_path.stem + '.pdf')
            if generated != pdf_path:
                generated.rename(pdf_path)
            return True, ''
        return False, result.stderr
    except Exception as e:
        return False, str(e)


def _find_chinese_font():
    for name in ('SimSun', 'Noto Sans CJK SC', 'WenQuanYi Micro Hei',
                 'Source Han Sans SC', 'Microsoft YaHei', 'PingFang SC'):
        try:
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase import pdfmetrics
            for path_str in [
                f'C:/Windows/Fonts/{name}.ttf',
                f'C:/Windows/Fonts/{name}.ttc',
                f'/usr/share/fonts/truetype/{name}.ttf',
                f'/System/Library/Fonts/{name}.ttf',
                f'/System/Library/Fonts/{name}.ttc',
            ]:
                p = Path(path_str)
                if p.exists():
                    pdfmetrics.registerFont(TTFont('ChineseFont', str(p)))
                    return 'ChineseFont'
        except Exception:
            continue
    return None


def _docx_to_pdf_simple(docx_path = None, pdf_path = None):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        font_name = _find_chinese_font()
        if not font_name:
            return False, '未找到中文字体'
        style = ParagraphStyle('body', fontName = font_name, fontSize = 12,
                               leading = 20, firstLineIndent = 24)
        doc = Document(docx_path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        report = SimpleDocTemplate(str(pdf_path), pagesize = A4,
                                   leftMargin = 2.5 * cm, rightMargin = 2.5 * cm,
                                   topMargin = 2.5 * cm, bottomMargin = 2.5 * cm)
        story = []
        for t in texts:
            story.append(Paragraph(t, style))
            story.append(Spacer(1, 6))
        report.build(story)
        return True, ''
    except Exception as e:
        return False, str(e)


def convert_project_word_to_pdf(project_dir = None, *, overwrite = False, progress = None):
    project_dir = Path(project_dir)
    results = []
    for docx_path in sorted(project_dir.rglob('*.docx')):
        pdf_path = docx_path.with_suffix('.pdf')
        if pdf_path.exists() and not overwrite:
            results.append((True, str(pdf_path), ''))
            continue
        ok, msg = _docx_to_pdf_word(docx_path, pdf_path)
        if not ok:
            ok, msg = _docx_to_pdf_office(docx_path, pdf_path)
        if not ok:
            ok, msg = _docx_to_pdf_simple(docx_path, pdf_path)
        results.append((ok, str(pdf_path), msg))
        if progress:
            try:
                progress(docx_path.name)
            except Exception:
                pass
    return results